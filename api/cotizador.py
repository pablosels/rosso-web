# -*- coding: utf-8 -*-
"""
Generador de cotizaciones de evento para ROSSO sobre el papel membretado oficial.

Uso:
    python cotizador.py evento.json

El JSON describe el evento (ver ejemplos en la carpeta eventos/). El documento
hereda encabezado (logo + datos de contacto), pie, fuentes y estilos de
plantilla/ROSSO_membretada.docx.
"""
import json
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

AQUI = os.path.dirname(os.path.abspath(__file__))
PLANTILLA = os.path.join(AQUI, "plantilla", "ROSSO_membretada.docx")

ROJO = RGBColor(0xB1, 0x07, 0x14)     # rojo del logo, muestreado del membrete
GRIS = RGBColor(0x55, 0x55, 0x55)
NEGRO = RGBColor(0x11, 0x11, 0x11)
FUENTE = "Helvetica Neue"

# etiquetas fijas del documento, por idioma
TEXTOS = {
    "es": {"kicker": "Cotización de evento", "concepto": "CONCEPTO",
           "importe": "IMPORTE", "total": "TOTAL",
           "ficha": [("Cliente", "cliente"), ("Fecha", "fecha"),
                     ("Horario", "horario"), ("Duración", "duracion"),
                     ("Invitados", "invitados"), ("Modalidad", "modalidad")]},
    "en": {"kicker": "Event proposal", "concepto": "ITEM",
           "importe": "AMOUNT", "total": "TOTAL",
           "ficha": [("Client", "cliente"), ("Date", "fecha"),
                     ("Time", "horario"), ("Duration", "duracion"),
                     ("Guests", "invitados"), ("Format", "modalidad")]},
}


def dinero(n):
    return "${:,.0f} MXN".format(round(float(n)))


def vaciar(doc):
    """Deja el documento sin cuerpo pero conserva header/footer/estilos/sectPr."""
    body = doc.element.body
    for hijo in list(body):
        if hijo.tag != qn("w:sectPr"):
            body.remove(hijo)


def _run(p, texto, size=10, bold=False, color=NEGRO, caps=False, espaciado=None):
    r = p.add_run(texto)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FUENTE)
    if espaciado is not None:                      # tracking, en puntos
        el = rpr.makeelement(qn("w:spacing"),
                             {qn("w:val"): str(int(espaciado * 20))})
        rpr.append(el)
    return r


def parrafo(doc, texto="", size=10, bold=False, color=NEGRO, antes=0, despues=4,
            align=None, caps=False, espaciado=None, interlineado=1.25):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(antes)
    pf.space_after = Pt(despues)
    pf.line_spacing = interlineado
    if align is not None:
        p.alignment = align
    if texto:
        _run(p, texto, size, bold, color, caps, espaciado)
    return p


def titulo_seccion(doc, texto):
    p = parrafo(doc, texto, size=9, bold=True, color=ROJO, antes=16,
                despues=6, caps=True, espaciado=1.4)
    p.paragraph_format.keep_with_next = True   # el titulo no se queda huerfano
    return p


def vineta(doc, texto, size=9.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.6)
    pf.first_line_indent = Cm(-0.35)
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.2
    _run(p, "— ", size, color=ROJO)
    _run(p, texto, size, color=NEGRO)
    return p


def _borde(celda, lados, color="D8D8D8", sz=4):
    tcPr = celda._tc.get_or_add_tcPr()
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for lado in lados:
        borders.append(borders.makeelement(qn("w:" + lado), {
            qn("w:val"): "single", qn("w:sz"): str(sz),
            qn("w:space"): "0", qn("w:color"): color}))
    tcPr.append(borders)


def _sombra(celda, color):
    tcPr = celda._tc.get_or_add_tcPr()
    tcPr.append(tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color}))


def tabla_conceptos(doc, filas, total=None, etiqueta_total="TOTAL", txt=None):
    """filas = [(concepto, importe, resalta_bool)]"""
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    anchos = (Cm(11.4), Cm(5.0))

    def _fila(a, b, bold=False, color=NEGRO, fondo=None, linea_arriba=False,
              size=9.5):
        r = t.add_row()
        datos = ((r.cells[0], a, WD_ALIGN_PARAGRAPH.LEFT),
                 (r.cells[1], b, WD_ALIGN_PARAGRAPH.RIGHT))
        for i, (celda, txt, al) in enumerate(datos):
            celda.width = anchos[i]
            p = celda.paragraphs[0]
            p.alignment = al
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            _run(p, txt, size, bold, color)
            lados = ["bottom"] + (["top"] if linea_arriba else [])
            _borde(celda, lados, color="B10714" if linea_arriba else "D8D8D8",
                   sz=8 if linea_arriba else 4)
            if fondo:
                _sombra(celda, fondo)

    txt = txt or TEXTOS["es"]
    _fila(txt["concepto"], txt["importe"], bold=True, color=GRIS, size=8)
    for concepto, importe, resalta in filas:
        _fila(concepto, importe, bold=resalta, color=NEGRO)
    if total is not None:
        _fila(etiqueta_total, total, bold=True, color=ROJO,
              fondo="F2F2F2", linea_arriba=True, size=10.5)
    # que la tabla no se parta entre paginas: el total nunca debe quedar solo
    for fila in t.rows[:-1]:
        for celda in fila.cells:
            for par in celda.paragraphs:
                par.paragraph_format.keep_with_next = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def generar(cfg, salida):
    doc = Document(PLANTILLA)
    vaciar(doc)
    txt = TEXTOS.get(cfg.get("idioma", "es"), TEXTOS["es"])

    parrafo(doc, cfg.get("kicker") or txt["kicker"], size=8.5, bold=True, color=ROJO,
            antes=28, despues=2, caps=True, espaciado=2.2)
    parrafo(doc, cfg["titulo"], size=21, bold=True, color=NEGRO,
            antes=0, despues=2, interlineado=1.05)
    if cfg.get("subtitulo"):
        parrafo(doc, cfg["subtitulo"], size=10, color=GRIS, antes=0, despues=14)

    ficha = [(k, cfg.get(v)) for k, v in txt["ficha"]]
    ficha = [(k, v) for k, v in ficha if v]
    t = doc.add_table(rows=0, cols=2)
    t.autofit = False
    for k, v in ficha:
        r = t.add_row()
        r.cells[0].width = Cm(3.6)
        r.cells[1].width = Cm(12.8)
        p0 = r.cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        _run(p0, k, 8, True, GRIS, caps=True, espaciado=0.8)
        p1 = r.cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        _run(p1, v, 10, False, NEGRO)
    parrafo(doc, "", despues=2)

    if cfg.get("intro"):
        parrafo(doc, cfg["intro"], size=10, antes=8, despues=4)

    for bloque in cfg.get("bloques", []):
        tipo = bloque.get("tipo", "lista")
        if bloque.get("salto"):
            doc.add_page_break()
        titulo_seccion(doc, bloque["titulo"])
        if bloque.get("nota"):
            nota = parrafo(doc, bloque["nota"], size=9.5, color=GRIS,
                           antes=0, despues=6)
            # solo las tablas: la nota nunca debe quedar sin su tabla
            nota.paragraph_format.keep_with_next = (tipo == "tabla")

        if tipo == "lista":
            for col in bloque.get("columnas", []):
                if col.get("titulo"):
                    parrafo(doc, col["titulo"], size=9.5, bold=True, antes=6,
                            despues=3)
                if col.get("nota"):
                    parrafo(doc, col["nota"], size=9, color=GRIS, antes=0,
                            despues=3)
                parrafo(doc, "   ·   ".join(col["items"]), size=9.5,
                        antes=0, despues=4, interlineado=1.3)

        elif tipo == "tabla":
            filas = [(f[0], f[1], bool(f[2]) if len(f) > 2 else False)
                     for f in bloque["filas"]]
            tabla_conceptos(doc, filas, bloque.get("total"),
                            bloque.get("etiqueta_total", txt["total"]), txt)

        elif tipo == "vinetas":
            for it in bloque["items"]:
                vineta(doc, it)

        elif tipo == "texto":
            for par in bloque["parrafos"]:
                parrafo(doc, par, size=9.5, antes=0, despues=5)

    if cfg.get("cierre"):
        parrafo(doc, cfg["cierre"], size=9, color=GRIS, antes=12, despues=0)

    doc.save(salida)
    return salida


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    ruta = sys.argv[1]
    cfg = json.load(open(ruta, encoding="utf-8"))
    nombre = cfg.get("archivo") or os.path.splitext(os.path.basename(ruta))[0] + ".docx"
    salida = os.path.join(AQUI, "salida", nombre)
    os.makedirs(os.path.dirname(salida), exist_ok=True)
    generar(cfg, salida)
    print("OK ->", salida)


if __name__ == "__main__":
    main()
